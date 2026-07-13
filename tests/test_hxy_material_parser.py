from __future__ import annotations

import importlib
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from packaging.requirements import Requirement


ASSIGNMENT_ID = "20000000-0000-0000-0000-000000000001"
MATERIAL_ID = "70000000-0000-0000-0000-000000000001"
JOB_ID = "80000000-0000-0000-0000-000000000001"
CREATED_AT = datetime(2026, 7, 10, 12, 0, tzinfo=timezone.utc)


def test_docx_test_dependency_is_explicitly_declared() -> None:
    requirements_path = Path(__file__).parents[1] / "apps" / "api" / "requirements.txt"
    dependencies = {
        Requirement(line).name.lower()
        for line in requirements_path.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.lstrip().startswith("#")
    }

    assert "python-docx" in dependencies


def test_markitdown_parser_extracts_text_and_docx_content(tmp_path: Path) -> None:
    module = importlib.import_module("apps.api.hxy_product.material_parser")
    text_path = tmp_path / "门店问题.txt"
    text_path.write_text("顾客到店后先询问当下状态。", encoding="utf-8")
    docx_path = tmp_path / "首店接待.docx"
    document = Document()
    document.add_heading("首店接待流程", level=1)
    document.add_paragraph("先问顾客状态，再介绍适合的服务。")
    document.save(docx_path)

    text_result = module.parse_with_markitdown(text_path)
    docx_result = module.parse_with_markitdown(docx_path)

    assert "顾客到店后先询问当下状态" in text_result.text_content
    assert "首店接待流程" in docx_result.text_content
    assert "先问顾客状态" in docx_result.text_content
    assert text_result.parser_name == docx_result.parser_name == "markitdown"
    assert text_result.parser_version
    assert text_result.warnings == ()


def test_markitdown_parser_rejects_empty_or_missing_sources(tmp_path: Path) -> None:
    module = importlib.import_module("apps.api.hxy_product.material_parser")
    empty = tmp_path / "empty.txt"
    empty.write_text("  \n", encoding="utf-8")

    with pytest.raises(module.MaterialParseError) as empty_error:
        module.parse_with_markitdown(empty)
    with pytest.raises(module.MaterialParseError) as missing_error:
        module.parse_with_markitdown(tmp_path / "missing.docx")

    assert empty_error.value.code == "empty_parse_output"
    assert empty_error.value.retryable is False
    assert missing_error.value.code == "source_missing"
    assert missing_error.value.retryable is False


def test_derived_artifact_keys_are_deterministic_and_not_user_controlled() -> None:
    module = importlib.import_module("apps.api.hxy_product.material_parser")

    keys = module.build_artifact_storage_keys(ASSIGNMENT_ID, MATERIAL_ID, JOB_ID)

    assert keys == {
        "normalized_markdown": (
            f"{ASSIGNMENT_ID}/{MATERIAL_ID}/derived/{JOB_ID}/normalized.md"
        ),
        "source_card": (
            f"{ASSIGNMENT_ID}/{MATERIAL_ID}/derived/{JOB_ID}/source-card.json"
        ),
    }
    assert ".." not in " ".join(keys.values())
    with pytest.raises(ValueError):
        module.build_artifact_storage_keys("../../outside", MATERIAL_ID, JOB_ID)


@pytest.mark.parametrize(
    ("origin", "authority"),
    [
        ("external", "reference"),
        ("internal", "working_material"),
        ("internal", "claimed_official"),
    ],
)
def test_source_card_preserves_claimed_authority_without_promoting_it(
    origin: str,
    authority: str,
) -> None:
    parser_module = importlib.import_module("apps.api.hxy_product.material_parser")
    source_module = importlib.import_module("apps.api.hxy_product.source_card")
    parsed = parser_module.MaterialParseResult(
        text_content="# 荷小悦首店资料\n\n先问顾客状态，再介绍服务。",
        title="荷小悦首店资料",
        parser_name="markitdown",
        parser_version="0.1.6",
        warnings=(),
    )
    material = {
        "material_id": MATERIAL_ID,
        "file_name": "首店资料.docx",
        "sha256": "a" * 64,
        "size_bytes": 1024,
        "understanding": {
            "document_type": "门店流程资料",
            "source_origin": origin,
            "authority_level": authority,
            "knowledge_scale": "micro",
            "domain": "operations",
            "summary": "首店接待流程资料。",
        },
    }

    card = source_module.build_source_card(material, parsed, created_at=CREATED_AT)

    assert card["version"] == "hxy-source-card.v1"
    assert card["source_id"] == f"material:{MATERIAL_ID}"
    assert card["source_hash"] == "a" * 64
    assert card["source_origin"] == origin
    assert card["authority_level"] == authority
    assert card["official_use_allowed"] is False
    assert "official_answer" in card["blocked_use"]
    assert "external_marketing" in card["blocked_use"]
    assert "financing_statement" in card["blocked_use"]
    assert "medical_claim" in card["blocked_use"]
    assert card["quality_signals"]["extracted_char_count"] == len(parsed.text_content)
    assert card["parser"]["name"] == "markitdown"
    if origin == "external":
        assert card["allowed_use"] == ["reference", "draft"]
